import docx

# create a new document
doc = docx.Document()

# create table with two columns and no rows
table = doc.add_table(rows=0, cols=2)

# add column headings

chapter = input("ENTER CHAPTER = ")
heading_cells = table.add_row().cells
heading_cells[0].text = 'S.NO'
heading_cells[1].text = chapter

# loop to add questions to table
i = 1
while True:
    question = input("Enter question (press Enter to stop): ")
    if not question:  # stop if user presses Enter without typing any text
        break
    # encode string using UTF-8 before adding to document
    question = question.encode('utf-8')
    # add S.NO and question to table
    cells = table.add_row().cells
    cells[0].text = str(i)
    cells[1].add_paragraph(question.decode('utf-8'))
    i += 1

# save the document
doc.save(f'{chapter}.docx')
